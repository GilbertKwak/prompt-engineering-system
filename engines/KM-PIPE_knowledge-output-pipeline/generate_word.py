#!/usr/bin/env python3
"""
generate_word.py — KM-PIPE v3.0 · A6 DocumentAgent
====================================================
Word(.docx) 보고서 자동 생성기
KM-PIPE Knowledge-to-Output Pipeline의 DocumentAgent 실체 구현

사용법:
    python generate_word.py --input data.json --output report.docx
    python generate_word.py --input data.json --output report.docx --template corporate

입력 JSON 형식 (KM-PIPE pipeline output):
    notion_data, word_doc_structure 필드 포함

Ref: T-09/KM-PIPE-MASTER-v3.0, engines/KM-PIPE_knowledge-output-pipeline/
"""

import argparse
import json
import os
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[ERROR] python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# ─────────────────────────────────────────────
# THEME CONFIGURATION
# ─────────────────────────────────────────────
THEMES = {
    "corporate": {
        "primary": RGBColor(0x01, 0x69, 0x6F),   # Hydra Teal (Nexus)
        "secondary": RGBColor(0x28, 0x25, 0x1D),  # Sylph Gray
        "accent": RGBColor(0x43, 0x7A, 0x22),     # Gridania Green
        "heading_font": "Calibri",
        "body_font": "Calibri",
        "base_size": 11,
    },
    "minimal": {
        "primary": RGBColor(0x1A, 0x1A, 0x1A),
        "secondary": RGBColor(0x55, 0x55, 0x55),
        "accent": RGBColor(0x00, 0x78, 0xD4),
        "heading_font": "Arial",
        "body_font": "Arial",
        "base_size": 11,
    },
    "research": {
        "primary": RGBColor(0x00, 0x33, 0x6B),
        "secondary": RGBColor(0x33, 0x33, 0x33),
        "accent": RGBColor(0xC8, 0x10, 0x26),
        "heading_font": "Times New Roman",
        "body_font": "Times New Roman",
        "base_size": 12,
    },
}


# ─────────────────────────────────────────────
# DOCUMENT BUILDER
# ─────────────────────────────────────────────
class KMPipeDocumentBuilder:
    """
    KM-PIPE A6 DocumentAgent — Word 보고서 빌더
    python-docx 기반, Notion JSON → .docx 완전 변환
    """

    def __init__(self, theme_name: str = "corporate"):
        self.doc = Document()
        self.theme = THEMES.get(theme_name, THEMES["corporate"])
        self._setup_document_styles()
        self.image_counter = 0

    def _setup_document_styles(self):
        """문서 기본 스타일 설정"""
        # 페이지 여백 설정
        section = self.doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        # Normal 스타일
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = self.theme["body_font"]
        font.size = Pt(self.theme["base_size"])
        font.color.rgb = self.theme["secondary"]

        # Heading 1
        h1 = self.doc.styles["Heading 1"]
        h1.font.name = self.theme["heading_font"]
        h1.font.size = Pt(20)
        h1.font.bold = True
        h1.font.color.rgb = self.theme["primary"]

        # Heading 2
        h2 = self.doc.styles["Heading 2"]
        h2.font.name = self.theme["heading_font"]
        h2.font.size = Pt(16)
        h2.font.bold = True
        h2.font.color.rgb = self.theme["primary"]

        # Heading 3
        h3 = self.doc.styles["Heading 3"]
        h3.font.name = self.theme["heading_font"]
        h3.font.size = Pt(13)
        h3.font.bold = True
        h3.font.color.rgb = self.theme["secondary"]

    def add_cover_page(self, title: str, subtitle: str = "", meta: dict = None):
        """표지 페이지 생성"""
        # 상단 여백
        for _ in range(4):
            self.doc.add_paragraph()

        # 제목
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.add_run(title)
        run.font.name = self.theme["heading_font"]
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = self.theme["primary"]

        # 부제목
        if subtitle:
            p_sub = self.doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_sub = p_sub.add_run(subtitle)
            run_sub.font.size = Pt(14)
            run_sub.font.color.rgb = self.theme["secondary"]
            run_sub.font.italic = True

        # 구분선
        for _ in range(2):
            self.doc.add_paragraph()
        p_line = self.doc.add_paragraph("─" * 60)
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 메타 정보
        if meta:
            for _ in range(2):
                self.doc.add_paragraph()
            for key, val in meta.items():
                p_meta = self.doc.add_paragraph()
                p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_k = p_meta.add_run(f"{key}: ")
                run_k.font.bold = True
                run_k.font.size = Pt(10)
                run_v = p_meta.add_run(str(val))
                run_v.font.size = Pt(10)

        # 페이지 나누기
        self.doc.add_page_break()

    def add_abstract(self, summary: str):
        """초록(Abstract) 섹션"""
        self.doc.add_heading("Abstract", level=1)
        p = self.doc.add_paragraph()
        p.style = self.doc.styles["Normal"]
        run = p.add_run(summary)
        run.font.italic = True
        run.font.size = Pt(11)
        # 들여쓰기
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.right_indent = Cm(1)
        self.doc.add_paragraph()

    def add_section(self, heading: str, content: str, level: int = 2):
        """섹션 추가 (heading + content)"""
        self.doc.add_heading(heading, level=level)
        if content:
            # 여러 문단 처리
            paragraphs = content.split("\n\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    p = self.doc.add_paragraph(para_text)
                    p.style = self.doc.styles["Normal"]
                    p.paragraph_format.space_after = Pt(6)

    def add_key_points(self, key_points: list):
        """핵심 포인트 번호 리스트"""
        self.doc.add_heading("Key Points", level=2)
        for i, point in enumerate(key_points, 1):
            p = self.doc.add_paragraph(style="List Number")
            run = p.add_run(str(point))
            run.font.size = Pt(self.theme["base_size"])
        self.doc.add_paragraph()

    def add_insights(self, insights: list):
        """인사이트 섹션 (강조 박스 스타일)"""
        self.doc.add_heading("Key Insights", level=2)
        for i, insight in enumerate(insights, 1):
            p = self.doc.add_paragraph()
            run_num = p.add_run(f"💡 {i}. ")
            run_num.font.bold = True
            run_num.font.color.rgb = self.theme["accent"]
            run_text = p.add_run(str(insight))
            run_text.font.size = Pt(self.theme["base_size"])
            p.paragraph_format.space_after = Pt(8)
        self.doc.add_paragraph()

    def add_table(self, caption: str, headers: list, rows: list):
        """자동 표 생성 (헤더 색상 강조)"""
        if caption:
            p_cap = self.doc.add_paragraph()
            run_cap = p_cap.add_run(f"Table: {caption}")
            run_cap.font.bold = True
            run_cap.font.size = Pt(10)
            run_cap.font.italic = True

        if not headers or not rows:
            return

        col_count = len(headers)
        table = self.doc.add_table(rows=1, cols=col_count)
        table.style = "Table Grid"

        # 헤더 행
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = str(header)
            # 헤더 배경색 설정
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "01696F")  # Hydra Teal
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)
            # 헤더 폰트 흰색
            for para in hdr_cells[i].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True

        # 데이터 행
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_val in enumerate(row_data):
                if i < col_count:
                    row_cells[i].text = str(cell_val)

        self.doc.add_paragraph()

    def add_image_placeholder(self, caption: str, chart_type: str = ""):
        """이미지 삽입 위치 플레이스홀더"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[📊 그림 {self.image_counter + 1}: {caption}")
        if chart_type:
            run.text += f" ({chart_type})"
        run.text += "]"
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.italic = True
        run.font.size = Pt(10)
        self.image_counter += 1
        self.doc.add_paragraph()

    def add_image(self, image_path: str, caption: str = "", width_cm: float = 14.0):
        """실제 이미지 파일 삽입"""
        if not os.path.exists(image_path):
            self.add_image_placeholder(caption or image_path)
            return
        try:
            self.doc.add_picture(image_path, width=Cm(width_cm))
            last_para = self.doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption:
                p_cap = self.doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_cap = p_cap.add_run(f"Figure {self.image_counter + 1}: {caption}")
                run_cap.font.size = Pt(9)
                run_cap.font.italic = True
                run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            self.image_counter += 1
        except Exception as e:
            print(f"[WARN] 이미지 삽입 실패 ({image_path}): {e}")
            self.add_image_placeholder(caption or image_path)

    def add_backlinks(self, backlinks: list):
        """참조/백링크 섹션"""
        if not backlinks:
            return
        self.doc.add_heading("References & Backlinks", level=2)
        for link in backlinks:
            p = self.doc.add_paragraph(style="List Bullet")
            run = p.add_run(str(link))
            run.font.size = Pt(10)
        self.doc.add_paragraph()

    def add_kg_metadata(self, kg_data: dict):
        """KG 메타데이터 푸터 섹션"""
        self.doc.add_paragraph()
        p_line = self.doc.add_paragraph("─" * 80)
        p_meta = self.doc.add_paragraph()
        run = p_meta.add_run(
            f"KG Node: {kg_data.get('new_node_id', 'N/A')}  │  "
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M KST')}  │  "
            f"Engine: KM-PIPE v3.0 · A6-DocumentAgent"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run.font.italic = True

    def save(self, output_path: str) -> str:
        """파일 저장"""
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output_path)
        size_kb = os.path.getsize(output_path) / 1024
        print(f"[OK] Word 문서 저장 완료: {output_path} ({size_kb:.1f} KB)")
        return output_path


# ─────────────────────────────────────────────
# PIPELINE RUNNER
# ─────────────────────────────────────────────
def build_from_km_pipe_output(data: dict, output_path: str, theme: str = "corporate") -> str:
    """
    KM-PIPE JSON output → Word 문서 완전 변환
    
    Args:
        data: KM-PIPE pipeline output JSON
        output_path: 저장 경로 (.docx)
        theme: corporate | minimal | research
    
    Returns:
        저장된 파일 경로
    """
    builder = KMPipeDocumentBuilder(theme_name=theme)
    notion = data.get("notion_data", {})
    word_struct = data.get("word_doc_structure", {})
    kg_data = data.get("kg_delta", {})
    viz = data.get("visualization", {})

    title = word_struct.get("title") or notion.get("title", "KM-PIPE Report")
    domain = notion.get("domain", "")
    kg_node = notion.get("kg_node_id", "")
    date_str = notion.get("created_at", datetime.now().strftime("%Y-%m-%d"))

    # 1. 표지
    builder.add_cover_page(
        title=title,
        subtitle=f"{domain.upper()} Domain Report" if domain else "",
        meta={
            "Date": date_str,
            "KG Node": kg_node or "TBD",
            "Status": notion.get("status", "✅ 운영 중"),
            "Engine": "KM-PIPE v3.0 · A6-DocumentAgent",
        },
    )

    # 2. 초록
    abstract = word_struct.get("abstract") or notion.get("summary", "")
    if abstract:
        builder.add_abstract(abstract)

    # 3. 핵심 포인트
    key_points = notion.get("key_points", [])
    if key_points:
        builder.add_key_points(key_points)

    # 4. 인사이트
    insights = notion.get("insights", [])
    if insights:
        builder.add_insights(insights)

    # 5. 본문 섹션
    sections = word_struct.get("sections", [])
    for sec in sections:
        heading = sec.get("heading", "")
        content = sec.get("content", "")
        level = sec.get("level", 2)
        if heading:
            builder.add_section(heading, content, level)

    # 6. 표
    tables = word_struct.get("tables", [])
    for tbl in tables:
        builder.add_table(
            caption=tbl.get("caption", ""),
            headers=tbl.get("headers", []),
            rows=tbl.get("rows", []),
        )

    # 7. 이미지 플레이스홀더 / 실제 이미지
    image_placeholders = word_struct.get("image_placeholders", [])
    chart_type = viz.get("chart_type", "")
    for i, ph in enumerate(image_placeholders):
        caption = ph.get("caption", f"Chart {i+1}")
        # 실제 차트 파일이 있으면 삽입, 없으면 플레이스홀더
        chart_path = ph.get("file_path", "")
        if chart_path and os.path.exists(chart_path):
            builder.add_image(chart_path, caption)
        else:
            builder.add_image_placeholder(caption, ph.get("chart_type", chart_type))

    # 8. 백링크/참조
    backlinks = notion.get("backlinks", []) + notion.get("wikilinks", [])
    builder.add_backlinks(backlinks)

    # 9. KG 메타데이터 푸터
    builder.add_kg_metadata(kg_data)

    return builder.save(output_path)


# ─────────────────────────────────────────────
# DEMO — 샘플 데이터로 테스트
# ─────────────────────────────────────────────
def run_demo(output_path: str = "reports/demo_report.docx"):
    """샘플 KM-PIPE JSON으로 Word 문서 생성 데모"""
    sample_data = {
        "pipeline_status": "success",
        "pe_score_final": 92,
        "action": "create",
        "notion_data": {
            "title": "HBM4 시장 분석 — 2026 Q2 업데이트",
            "summary": "SK하이닉스의 HBM4 양산 일정이 앞당겨지며 AI 인프라 수요를 견인하고 있습니다."
                       "엔비디아 GB300 플랫폼과의 소켓 호환성이 핵심 경쟁 요소로 부상했습니다.",
            "key_points": [
                "SK하이닉스 HBM4 양산 2026 Q3 확정",
                "삼성전자 HBM4E 수율 개선으로 점유율 회복 중",
                "마이크론 HBM3E Gen2 NVIDIA 공급 확대",
                "CoWoS 패키징 병목이 HBM 공급 제약의 주요 인자",
                "중국 HBM 자체 개발 — ChangXin Memory 2027 목표",
            ],
            "insights": [
                "HBM4 전환기에 수율 우위 기업이 ASP 프리미엄 2~3배 수취 가능",
                "CoWoS 2.5D 패키징 투자 확대 기업(TSMC, Amkor)에 간접 수혜 기회 존재",
                "중국 자체 HBM 개발 실패시 2028년까지 SK하이닉스 독점 구조 유지 전망",
            ],
            "backlinks": ["[[C-37 AI Ecosystem Intelligence]]", "[[C-38 PE-INTEL]]", "[[T-09 Mother Page]]"],
            "domain": "semiconductor",
            "kg_node_id": "SEM-HBM4-20260523",
            "status": "✅ 운영 중",
            "created_at": "2026-05-23",
        },
        "word_doc_structure": {
            "title": "HBM4 시장 분석 — 2026 Q2 업데이트",
            "abstract": "본 보고서는 HBM4 메모리 시장의 2026년 2분기 현황을 분석합니다. "
                        "주요 공급업체(SK하이닉스, 삼성전자, 마이크론)의 양산 일정 및 기술 로드맵을 검토하고, "
                        "AI 인프라 수요와의 연계성을 바탕으로 투자 시사점을 도출합니다.",
            "sections": [
                {"heading": "시장 현황", "level": 2,
                 "content": "HBM 시장은 2026년에 AI 가속기 수요 폭증으로 전년 대비 65% 성장이 예상됩니다. "
                            "SK하이닉스가 약 52%의 시장 점유율을 유지하며 선두를 달리고 있습니다."},
                {"heading": "경쟁 구도 분석", "level": 2,
                 "content": "3사 경쟁 구도에서 기술력과 수율이 핵심 차별화 요소입니다. "
                            "HBM4는 12-high 스택으로 이전 세대 대비 대역폭 50% 향상을 달성했습니다."},
                {"heading": "투자 시사점", "level": 2,
                 "content": "HBM4 양산 안정화 이후 ASP 하락 압력이 예상되나, "
                            "수요 증가세가 이를 상회할 것으로 전망됩니다."},
            ],
            "tables": [
                {
                    "caption": "HBM 공급업체 비교 (2026 Q2)",
                    "headers": ["업체", "제품", "양산 시기", "시장점유율", "주요 고객"],
                    "rows": [
                        ["SK하이닉스", "HBM4", "2026 Q3", "52%", "NVIDIA, AMD"],
                        ["삼성전자", "HBM4E", "2026 Q4", "33%", "NVIDIA, Google"],
                        ["마이크론", "HBM3E Gen2", "2026 Q2 (출하중)", "15%", "NVIDIA"],
                    ],
                }
            ],
            "image_placeholders": [
                {"position": "after_section_1", "caption": "HBM 시장 점유율 추이", "chart_type": "line"},
                {"position": "after_section_2", "caption": "공급업체 경쟁 포지셔닝", "chart_type": "scatter"},
            ],
        },
        "kg_delta": {
            "new_node_id": "SEM-HBM4-20260523",
            "update_command": "python automation/kg_updater.py --add-node SEM-HBM4-20260523",
        },
        "visualization": {"chart_type": "line"},
    }
    print("[DEMO] 샘플 데이터로 Word 문서 생성 중...")
    result = build_from_km_pipe_output(sample_data, output_path, theme="corporate")
    return result


# ─────────────────────────────────────────────
# CLI ENTRY POINT
# ─────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description="KM-PIPE v3.0 · A6 DocumentAgent — Word 보고서 자동 생성",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
예시:
  python generate_word.py --demo
  python generate_word.py --input reports/km_pipe_output.json --output reports/report.docx
  python generate_word.py --input data.json --output report.docx --theme research
        """,
    )
    parser.add_argument("--input", "-i", help="KM-PIPE JSON output 파일 경로")
    parser.add_argument("--output", "-o", default="reports/km_pipe_report.docx", help="출력 .docx 경로")
    parser.add_argument("--theme", "-t", choices=["corporate", "minimal", "research"],
                        default="corporate", help="문서 테마")
    parser.add_argument("--demo", action="store_true", help="샘플 데이터로 데모 실행")

    args = parser.parse_args()

    if args.demo:
        result = run_demo(args.output)
        print(f"[DEMO COMPLETE] {result}")
        return

    if not args.input:
        parser.error("--input 또는 --demo 옵션 중 하나가 필요합니다.")

    if not os.path.exists(args.input):
        print(f"[ERROR] 입력 파일 없음: {args.input}")
        sys.exit(1)

    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    result = build_from_km_pipe_output(data, args.output, theme=args.theme)
    print(f"[DONE] {result}")


if __name__ == "__main__":
    main()
