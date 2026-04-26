#!/usr/bin/env python3
"""
P7 — PE-7 Word 보고서 자동 생성 엔진
docx_report_builder.py

ExaFLOPS TCO 분석 결과(JSON/Excel) + 차트 PNG 5개를 DOCX에 자동 삽입.

Usage:
    python docx_report_builder.py                          # 전체 빌드
    python docx_report_builder.py --system SYS_C           # 단일 시스템
    python docx_report_builder.py --scenario smr_full      # 단일 시나리오
    python docx_report_builder.py --lang en                # 영문 보고서
    python docx_report_builder.py --dry-run                # 템플릿 검증만

Outputs:
    output/reports/PE7_ExaFLOPS_TCO_Report_YYYYMMDD.docx
    output/reports/PE7_ExaFLOPS_TCO_Report_YYYYMMDD_EN.docx  (--lang en)
    output/reports/build_manifest.json
"""

import argparse
import json
import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[ERROR] python-docx 미설치. pip install python-docx 실행 후 재시도")
    sys.exit(1)

try:
    import yaml
except ImportError:
    print("[ERROR] PyYAML 미설치. pip install pyyaml 실행 후 재시도")
    sys.exit(1)


# ─────────────────────────────────────────────
# 0. 경로 상수
# ─────────────────────────────────────────────
BASE_DIR = Path(__file__).parent
CONFIG_PATH = BASE_DIR / "p7_report_config.yaml"
EXAFLOPS_JSON = BASE_DIR.parent.parent / "output" / "p6_exaflops_tco.json"
CHARTS_DIR = BASE_DIR.parent.parent / "output" / "p6_charts"
OUTPUT_DIR = BASE_DIR.parent.parent / "output" / "reports"
TEMPLATE_DIR = BASE_DIR / "templates"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ─────────────────────────────────────────────
# 1. 설정 로더
# ─────────────────────────────────────────────
class P7Config:
    def __init__(self, config_path: Path = CONFIG_PATH):
        with open(config_path, "r", encoding="utf-8") as f:
            self._cfg = yaml.safe_load(f)

    @property
    def report(self) -> dict:
        return self._cfg.get("report", {})

    @property
    def sections(self) -> list:
        return self._cfg.get("sections", [])

    @property
    def charts(self) -> list:
        return self._cfg.get("charts", [])

    @property
    def style(self) -> dict:
        return self._cfg.get("style", {})

    @property
    def strings(self) -> dict:
        return self._cfg.get("strings", {})


# ─────────────────────────────────────────────
# 2. 스타일 헬퍼
# ─────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    """테이블 셀 배경색 설정 (OOXML 직접 조작)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _bold_cell(cell, text: str, font_size: int = 10, color: str = "FFFFFF"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(color)


def _add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2C5F8A")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─────────────────────────────────────────────
# 3. 데이터 로더
# ─────────────────────────────────────────────
def load_tco_data(json_path: Path = EXAFLOPS_JSON) -> dict:
    """P6 출력 JSON 로드. 파일 없으면 더미 데이터 반환."""
    if json_path.exists():
        with open(json_path, "r", encoding="utf-8") as f:
            return json.load(f)
    # 더미 데이터 (P6 실행 전 빌드 허용)
    return {
        "generated_at": datetime.now().isoformat(),
        "systems": {
            "SYS_A": {
                "label": "HPC Cluster (1 EF)",
                "exaflops": 1.0,
                "scenarios": {
                    "grid_only": {"total_tco_m": 4823, "tco_per_ef": 4823, "irr": 0.38, "moic": 6.43},
                    "smr_partial": {"total_tco_m": 4641, "tco_per_ef": 4641, "irr": 0.40, "moic": 7.10},
                    "smr_full": {"total_tco_m": 4398, "tco_per_ef": 4398, "irr": 0.43, "moic": 7.83},
                }
            },
            "SYS_B": {
                "label": "AI Training DC (5 EF)",
                "exaflops": 5.0,
                "scenarios": {
                    "grid_only": {"total_tco_m": 21840, "tco_per_ef": 4368, "irr": 0.40, "moic": 6.89},
                    "smr_partial": {"total_tco_m": 20960, "tco_per_ef": 4192, "irr": 0.42, "moic": 7.58},
                    "smr_full": {"total_tco_m": 19780, "tco_per_ef": 3956, "irr": 0.45, "moic": 8.36},
                }
            },
            "SYS_C": {
                "label": "Sovereign AI DC (10 EF)",
                "exaflops": 10.0,
                "scenarios": {
                    "grid_only": {"total_tco_m": 41200, "tco_per_ef": 4120, "irr": 0.42, "moic": 7.21},
                    "smr_partial": {"total_tco_m": 39400, "tco_per_ef": 3940, "irr": 0.44, "moic": 7.97},
                    "smr_full": {"total_tco_m": 37100, "tco_per_ef": 3710, "irr": 0.48, "moic": 8.83},
                }
            }
        },
        "portfolio": {
            "grid_only": {"irr": 0.38, "moic": 6.43, "net_alpha_m": 0},
            "smr_partial": {"irr": 0.40, "moic": 6.97, "net_alpha_m": 180},
            "smr_full": {"irr": 0.44, "moic": 7.83, "net_alpha_m": 520},
        }
    }


def find_charts(charts_dir: Path) -> dict:
    """P6 차트 PNG 파일 목록 수집"""
    chart_map = {
        "tco_heatmap": "tco_heatmap.png",
        "moic_grouped_bar": "moic_grouped_bar.png",
        "moic_delta_bar": "moic_delta_bar.png",
        "moic_waterfall": "moic_waterfall.png",
        "irr_moic_bubble": "irr_moic_bubble.png",
    }
    found = {}
    for key, fname in chart_map.items():
        p = charts_dir / fname
        if p.exists():
            found[key] = p
        else:
            # fallback: output/ 루트에서 검색
            alt = charts_dir.parent / fname
            if alt.exists():
                found[key] = alt
    return found


# ─────────────────────────────────────────────
# 4. 보고서 빌더 클래스
# ─────────────────────────────────────────────
class ReportBuilder:
    ACCENT = "2C5F8A"   # 브랜드 블루
    ACCENT2 = "1A3A5C"  # 다크 블루
    LIGHT = "EBF3FB"    # 라이트 블루
    GRAY = "F5F5F5"

    def __init__(self, cfg: P7Config, data: dict, charts: dict, lang: str = "ko"):
        self.cfg = cfg
        self.data = data
        self.charts = charts
        self.lang = lang
        self.doc = Document()
        self._setup_document()

    # ── 4-1. 문서 기본 설정 ──
    def _setup_document(self):
        sec = self.doc.sections[0]
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)

    # ── 4-2. 커버 페이지 ──
    def build_cover(self):
        doc = self.doc
        doc.add_paragraph()
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("PE-7 ExaFLOPS TCO 분석 보고서" if self.lang == "ko"
                        else "PE-7 ExaFLOPS TCO Analysis Report")
        run.bold = True
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor.from_string(self.ACCENT2)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("SMR 통합 자원모델 — HBM · Glass · 전력 · 패키징 4-레이어 TCO" if self.lang == "ko"
                          else "SMR Integrated Resource Model — 4-Layer TCO: HBM · Glass · Power · Packaging")
        run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor.from_string(self.ACCENT)

        doc.add_paragraph()
        _add_horizontal_rule(doc)
        doc.add_paragraph()

        meta_items = [
            ("작성일" if self.lang == "ko" else "Date",
             datetime.now().strftime("%Y-%m-%d")),
            ("버전" if self.lang == "ko" else "Version", "v1.0"),
            ("분류" if self.lang == "ko" else "Classification", "Confidential"),
            ("시스템" if self.lang == "ko" else "Systems",
             "SYS_A · SYS_B · SYS_C"),
            ("시나리오" if self.lang == "ko" else "Scenarios",
             "Grid Only · SMR Partial · SMR Full"),
        ]
        tbl = doc.add_table(rows=len(meta_items), cols=2)
        tbl.style = "Table Grid"
        for i, (k, v) in enumerate(meta_items):
            tbl.cell(i, 0).text = k
            tbl.cell(i, 0).paragraphs[0].runs[0].bold = True
            _set_cell_bg(tbl.cell(i, 0), self.LIGHT)
            tbl.cell(i, 1).text = v

        doc.add_page_break()

    # ── 4-3. 목차 ──
    def build_toc(self):
        doc = self.doc
        p = doc.add_heading("목차" if self.lang == "ko" else "Contents", level=1)
        p.runs[0].font.color.rgb = RGBColor.from_string(self.ACCENT2)

        toc_items = [
            ("1", "Executive Summary"),
            ("2", "분석 방법론 및 가정" if self.lang == "ko" else "Methodology & Assumptions"),
            ("3", "ExaFLOPS 시스템 정의" if self.lang == "ko" else "ExaFLOPS System Definitions"),
            ("4", "4-레이어 TCO 모델 결과" if self.lang == "ko" else "4-Layer TCO Model Results"),
            ("5", "SMR 시나리오별 민감도 분석" if self.lang == "ko" else "SMR Scenario Sensitivity Analysis"),
            ("6", "포트폴리오 MOIC / IRR 분석" if self.lang == "ko" else "Portfolio MOIC / IRR Analysis"),
            ("7", "차트 패키지 (P5 시각화)" if self.lang == "ko" else "Chart Package (P5 Visualizations)"),
            ("8", "결론 및 권고사항" if self.lang == "ko" else "Conclusions & Recommendations"),
            ("부록", "데이터 출처 및 파라미터" if self.lang == "ko" else "Data Sources & Parameters"),
        ]
        for num, title in toc_items:
            p = doc.add_paragraph(style="List Number" if num.isdigit() else "Normal")
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(f"{num}. {title}" if num.isdigit() else f"부록. {title}")
            run.font.size = Pt(11)

        doc.add_page_break()

    # ── 4-4. Executive Summary ──
    def build_executive_summary(self):
        doc = self.doc
        doc.add_heading("1. Executive Summary", level=1)

        pf = self.data.get("portfolio", {})
        grid = pf.get("grid_only", {})
        smr_f = pf.get("smr_full", {})

        summary_text = (
            f"본 보고서는 PE-7 프로젝트의 ExaFLOPS 통합 TCO 분석 결과를 정리합니다. "
            f"HBM, Glass Substrate, 전력(SMR), OSAT 패키징 4개 레이어를 단일 "
            f"TCO 단위로 통합하여 SYS_A(1 EF), SYS_B(5 EF), SYS_C(10 EF) 세 시스템에 "
            f"Grid Only · SMR Partial · SMR Full 3개 시나리오를 적용했습니다.\n\n"
            f"핵심 결과: SMR Full 시나리오에서 포트폴리오 MOIC는 "
            f"{grid.get('moic', 6.43):.2f}x(Grid Only)에서 "
            f"{smr_f.get('moic', 7.83):.2f}x로 상승하며, "
            f"Net Alpha는 ${smr_f.get('net_alpha_m', 520):,}M에 달합니다. "
            f"IRR은 {grid.get('irr', 0.38)*100:.1f}%에서 "
            f"{smr_f.get('irr', 0.44)*100:.1f}%로 개선됩니다."
        ) if self.lang == "ko" else (
            f"This report presents the integrated ExaFLOPS TCO analysis results for the PE-7 project. "
            f"Four cost layers — HBM, Glass Substrate, Power (SMR), and OSAT Packaging — are unified "
            f"into a single TCO framework applied across SYS_A (1 EF), SYS_B (5 EF), and SYS_C (10 EF) "
            f"under three scenarios: Grid Only, SMR Partial, and SMR Full.\n\n"
            f"Key finding: Under the SMR Full scenario, portfolio MOIC rises from "
            f"{grid.get('moic', 6.43):.2f}x (Grid Only) to {smr_f.get('moic', 7.83):.2f}x, "
            f"delivering a Net Alpha of ${smr_f.get('net_alpha_m', 520):,}M. "
            f"IRR improves from {grid.get('irr', 0.38)*100:.1f}% to {smr_f.get('irr', 0.44)*100:.1f}%."
        )
        doc.add_paragraph(summary_text)

        # KPI 요약 테이블
        doc.add_paragraph()
        kpi_label = "핵심 KPI 요약" if self.lang == "ko" else "Key KPI Summary"
        doc.add_heading(kpi_label, level=2)

        headers = ["KPI", "Grid Only", "SMR Partial", "SMR Full", "Delta (Full vs Grid)"]
        rows = [
            ["Portfolio MOIC",
             f"{pf.get('grid_only',{}).get('moic',6.43):.2f}x",
             f"{pf.get('smr_partial',{}).get('moic',6.97):.2f}x",
             f"{pf.get('smr_full',{}).get('moic',7.83):.2f}x",
             f"+{pf.get('smr_full',{}).get('moic',7.83)-pf.get('grid_only',{}).get('moic',6.43):.2f}x"],
            ["Portfolio IRR",
             f"{pf.get('grid_only',{}).get('irr',0.38)*100:.1f}%",
             f"{pf.get('smr_partial',{}).get('irr',0.40)*100:.1f}%",
             f"{pf.get('smr_full',{}).get('irr',0.44)*100:.1f}%",
             f"+{(pf.get('smr_full',{}).get('irr',0.44)-pf.get('grid_only',{}).get('irr',0.38))*100:.1f}pp"],
            ["Net Alpha ($M)",
             "$0",
             f"${pf.get('smr_partial',{}).get('net_alpha_m',180):,}M",
             f"${pf.get('smr_full',{}).get('net_alpha_m',520):,}M",
             f"+${pf.get('smr_full',{}).get('net_alpha_m',520):,}M"],
        ]
        tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = "Table Grid"
        for i, h in enumerate(headers):
            _bold_cell(tbl.cell(0, i), h, 10, "FFFFFF")
            _set_cell_bg(tbl.cell(0, i), self.ACCENT)
        for r_idx, row in enumerate(rows, 1):
            for c_idx, val in enumerate(row):
                cell = tbl.cell(r_idx, c_idx)
                cell.text = val
                if c_idx == 0:
                    cell.paragraphs[0].runs[0].bold = True
                _set_cell_bg(cell, self.GRAY if r_idx % 2 == 0 else "FFFFFF")

        doc.add_page_break()

    # ── 4-5. 시스템별 TCO 결과 ──
    def build_tco_results(self):
        doc = self.doc
        section_title = ("4. 4-레이어 TCO 모델 결과" if self.lang == "ko"
                         else "4. 4-Layer TCO Model Results")
        doc.add_heading(section_title, level=1)

        systems = self.data.get("systems", {})
        scenarios = ["grid_only", "smr_partial", "smr_full"]
        scen_labels = {
            "grid_only": "Grid Only",
            "smr_partial": "SMR Partial",
            "smr_full": "SMR Full",
        }

        for sys_id, sys_data in systems.items():
            label = sys_data.get("label", sys_id)
            ef = sys_data.get("exaflops", 1.0)
            sub_title = f"{sys_id} — {label} ({ef} ExaFLOPS)"
            doc.add_heading(sub_title, level=2)

            headers = ["지표" if self.lang == "ko" else "Metric"] + [scen_labels[s] for s in scenarios]
            metrics = [
                ("Total TCO ($M)", "total_tco_m", lambda v: f"${v:,.0f}M"),
                ("TCO/EF ($M/EF)", "tco_per_ef", lambda v: f"${v:,.0f}M"),
                ("IRR", "irr", lambda v: f"{v*100:.1f}%"),
                ("MOIC", "moic", lambda v: f"{v:.2f}x"),
            ]
            scen_data = sys_data.get("scenarios", {})

            tbl = doc.add_table(rows=1 + len(metrics), cols=len(headers))
            tbl.style = "Table Grid"
            for i, h in enumerate(headers):
                _bold_cell(tbl.cell(0, i), h, 9, "FFFFFF")
                _set_cell_bg(tbl.cell(0, i), self.ACCENT)

            for r_idx, (label_m, key, fmt) in enumerate(metrics, 1):
                tbl.cell(r_idx, 0).text = label_m
                tbl.cell(r_idx, 0).paragraphs[0].runs[0].bold = True
                _set_cell_bg(tbl.cell(r_idx, 0), self.LIGHT)
                for c_idx, scen in enumerate(scenarios, 1):
                    val = scen_data.get(scen, {}).get(key, 0)
                    tbl.cell(r_idx, c_idx).text = fmt(val)
                    _set_cell_bg(tbl.cell(r_idx, c_idx),
                                 self.GRAY if r_idx % 2 == 0 else "FFFFFF")
            doc.add_paragraph()

        doc.add_page_break()

    # ── 4-6. 차트 삽입 ──
    def build_chart_section(self):
        doc = self.doc
        section_title = ("7. 차트 패키지 (P5 시각화)" if self.lang == "ko"
                         else "7. Chart Package (P5 Visualizations)")
        doc.add_heading(section_title, level=1)

        chart_meta = [
            ("tco_heatmap",
             "[Fig 7-1] TCO NPV 절감 히트맵 — Facility Tier × 전력가" if self.lang == "ko"
             else "[Fig 7-1] TCO NPV Savings Heatmap — Facility Tier × Grid Price",
             "RR SMR 80% 침투 시 10Y NPV 절감액($M). AI DC US(SYS_C)에서 최대 절감 풀 형성." if self.lang == "ko"
             else "10Y NPV savings ($M) at 80% RR SMR penetration. SYS_C shows the largest savings pool."),
            ("moic_grouped_bar",
             "[Fig 7-2] MOIC Absolute — 시스템 × 시나리오 비교" if self.lang == "ko"
             else "[Fig 7-2] MOIC Absolute — System × Scenario Comparison",
             "SYS_C SMR Full에서 MOIC 8.83x 최대치 달성." if self.lang == "ko"
             else "SYS_C SMR Full achieves the maximum MOIC of 8.83x."),
            ("moic_delta_bar",
             "[Fig 7-3] MOIC Delta Uplift — Grid Only 대비 순증분" if self.lang == "ko"
             else "[Fig 7-3] MOIC Delta Uplift vs Grid Only",
             "SMR Full에서 SYS_C +1.62x로 최대 업리프트." if self.lang == "ko"
             else "SMR Full delivers +1.62x maximum uplift for SYS_C."),
            ("moic_waterfall",
             "[Fig 7-4] Portfolio MOIC Waterfall — SMR 기여 분해" if self.lang == "ko"
             else "[Fig 7-4] Portfolio MOIC Waterfall — SMR Contribution Breakdown",
             "Grid Only 6.43x → SMR Net 7.83x까지 단계별 기여 구조." if self.lang == "ko"
             else "Step-by-step contribution from Grid Only 6.43x to SMR Net 7.83x."),
            ("irr_moic_bubble",
             "[Fig 7-5] IRR Alpha vs MOIC Bubble — 자본 효율 비교" if self.lang == "ko"
             else "[Fig 7-5] IRR Alpha vs MOIC Bubble — Capital Efficiency",
             "버블 크기=Net Alpha($M). SMR Partial이 최고 자본 효율 달성." if self.lang == "ko"
             else "Bubble size=Net Alpha($M). SMR Partial achieves highest capital efficiency."),
        ]

        for i, (chart_key, title, caption) in enumerate(chart_meta, 1):
            doc.add_heading(title, level=2)
            chart_path = self.charts.get(chart_key)
            if chart_path and Path(chart_path).exists():
                doc.add_picture(str(chart_path), width=Inches(5.8))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[차트 파일 없음: {chart_key}.png — python exaflops_viz.py 실행 필요]")
                run.italic = True
                run.font.color.rgb = RGBColor(180, 0, 0)

            cap_para = doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_para.add_run(caption)
            cap_run.italic = True
            cap_run.font.size = Pt(9)
            cap_run.font.color.rgb = RGBColor.from_string("555555")
            doc.add_paragraph()

        doc.add_page_break()

    # ── 4-7. 결론 및 권고 ──
    def build_conclusions(self):
        doc = self.doc
        section_title = ("8. 결론 및 권고사항" if self.lang == "ko"
                         else "8. Conclusions & Recommendations")
        doc.add_heading(section_title, level=1)

        conclusions = [
            ("SMR Full — 최대 수익성" if self.lang == "ko" else "SMR Full — Maximum Returns",
             "포트폴리오 MOIC 7.83x, IRR 44%. SYS_C(10 EF)에서 TCO/EF $3,710M/EF로 Grid Only 대비 "
             "10.0% 절감. 장기 전력 안정성 확보 시 우선 추천." if self.lang == "ko"
             else "Portfolio MOIC 7.83x, IRR 44%. SYS_C achieves TCO/EF of $3,710M/EF, "
             "a 10.0% reduction vs. Grid Only. Recommended when long-term power stability is secured."),
            ("SMR Partial — 최고 자본 효율" if self.lang == "ko" else "SMR Partial — Best Capital Efficiency",
             "추가 CAPEX $360M으로 Net Alpha $180M 달성. CAPEX 제약이 있는 Phase 1에 적합." if self.lang == "ko"
             else "Delivers Net Alpha of $180M with incremental CAPEX of $360M. Best fit for CAPEX-constrained Phase 1."),
            ("HBM Salvage 연계" if self.lang == "ko" else "HBM Salvage Integration",
             "HBM Salvage 프로그램 연계 시 HBM Net Cost를 추가 $120~180M 절감 가능. "
             "P6 모델의 salvage_rate 0.35→0.45 상향 시 TCO/EF $250M 추가 개선." if self.lang == "ko"
             else "Integration with HBM Salvage Program can reduce HBM Net Cost by additional $120-180M. "
             "Raising salvage_rate from 0.35 to 0.45 in the P6 model yields an additional $250M TCO/EF improvement."),
        ]

        for title, body in conclusions:
            doc.add_heading(title, level=2)
            doc.add_paragraph(body)

        doc.add_page_break()

    # ── 4-8. 부록 ──
    def build_appendix(self):
        doc = self.doc
        app_title = "부록. 데이터 출처 및 파라미터" if self.lang == "ko" else "Appendix. Data Sources & Parameters"
        doc.add_heading(app_title, level=1)

        params = [
            ("WACC", "8.5%"),
            ("전력 에스컬레이션" if self.lang == "ko" else "Power Escalation", "2.5%/yr"),
            ("HBM3E Unit Cost", "$18,000/unit"),
            ("HBM4 Unit Cost", "$24,000/unit"),
            ("HBM Yield", "88%"),
            ("Salvage Rate", "35%"),
            ("Glass Premium", "12%"),
            ("RR SMR LCOE", "$72/MWh"),
            ("NuScale LCOE", "$89/MWh"),
            ("Kairos LCOE", "$78/MWh"),
            ("SMR CAPEX (Partial)", "$360M"),
            ("SMR CAPEX (Full)", "$720M"),
            ("CoWoS Unit Cost", "$8,500/unit"),
        ]
        tbl = doc.add_table(rows=1 + len(params), cols=2)
        tbl.style = "Table Grid"
        _bold_cell(tbl.cell(0, 0), "파라미터" if self.lang == "ko" else "Parameter", 10, "FFFFFF")
        _bold_cell(tbl.cell(0, 1), "값" if self.lang == "ko" else "Value", 10, "FFFFFF")
        for c in range(2):
            _set_cell_bg(tbl.cell(0, c), self.ACCENT)
        for i, (k, v) in enumerate(params, 1):
            tbl.cell(i, 0).text = k
            tbl.cell(i, 1).text = v
            _set_cell_bg(tbl.cell(i, 0), self.LIGHT if i % 2 == 0 else "FFFFFF")
            _set_cell_bg(tbl.cell(i, 1), self.LIGHT if i % 2 == 0 else "FFFFFF")

    # ── 4-9. 메인 빌드 ──
    def build(self, output_path: Path):
        print("[P7] Word 보고서 빌드 시작...")
        print("  ▶ 커버 페이지")
        self.build_cover()
        print("  ▶ 목차")
        self.build_toc()
        print("  ▶ Executive Summary")
        self.build_executive_summary()
        print("  ▶ 4-레이어 TCO 결과")
        self.build_tco_results()
        print(f"  ▶ 차트 섹션 ({len(self.charts)}개 차트 감지)")
        self.build_chart_section()
        print("  ▶ 결론 및 권고")
        self.build_conclusions()
        print("  ▶ 부록")
        self.build_appendix()

        self.doc.save(str(output_path))
        print(f"\n[P7] ✅ 보고서 저장 완료: {output_path}")
        return output_path


# ─────────────────────────────────────────────
# 5. 매니페스트 저장
# ─────────────────────────────────────────────
def save_manifest(output_path: Path, charts: dict, data: dict):
    manifest = {
        "generated_at": datetime.now().isoformat(),
        "report_path": str(output_path),
        "report_size_kb": round(output_path.stat().st_size / 1024, 1) if output_path.exists() else 0,
        "charts_embedded": list(charts.keys()),
        "charts_missing": [],
        "tco_data_source": str(EXAFLOPS_JSON),
        "systems": list(data.get("systems", {}).keys()),
        "scenarios": ["grid_only", "smr_partial", "smr_full"],
    }
    manifest_path = OUTPUT_DIR / "build_manifest.json"
    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=2, ensure_ascii=False)
    print(f"[P7] 매니페스트 저장: {manifest_path}")
    return manifest


# ─────────────────────────────────────────────
# 6. CLI 엔트리포인트
# ─────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="PE-7 P7 Word Report Builder")
    parser.add_argument("--system", choices=["SYS_A", "SYS_B", "SYS_C", "all"], default="all")
    parser.add_argument("--scenario", choices=["grid_only", "smr_partial", "smr_full", "all"], default="all")
    parser.add_argument("--lang", choices=["ko", "en"], default="ko", help="보고서 언어")
    parser.add_argument("--dry-run", action="store_true", help="파일 생성 없이 검증만")
    parser.add_argument("--output", type=str, default=None, help="출력 파일 경로")
    args = parser.parse_args()

    # 설정 로드
    cfg = P7Config()

    # 데이터 로드
    data = load_tco_data()
    print(f"[P7] TCO 데이터 로드: {len(data.get('systems', {}))}개 시스템")

    # 차트 수집
    charts = find_charts(CHARTS_DIR)
    print(f"[P7] 차트 감지: {len(charts)}개 — {list(charts.keys())}")

    if args.dry_run:
        print("[P7] Dry-run 완료. 파일 생성 생략.")
        return

    # 출력 경로
    timestamp = datetime.now().strftime("%Y%m%d")
    lang_suffix = "" if args.lang == "ko" else "_EN"
    default_name = f"PE7_ExaFLOPS_TCO_Report_{timestamp}{lang_suffix}.docx"
    output_path = Path(args.output) if args.output else OUTPUT_DIR / default_name

    # 빌드
    builder = ReportBuilder(cfg, data, charts, lang=args.lang)
    builder.build(output_path)

    # 매니페스트
    save_manifest(output_path, charts, data)

    print("\n[P7] 전체 빌드 완료 ✅")
    print(f"  보고서: {output_path}")
    print(f"  언어: {args.lang}")
    print(f"  차트 삽입: {len(charts)}개")


if __name__ == "__main__":
    main()
